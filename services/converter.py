"""
Сервис конвертации изображений и документов в различные форматы.
Использует reportlab для PDF, Pillow для изображений, PyMuPDF для PDF, python-docx для DOCX.
"""
import io
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from pdf2docx import Converter as PDF2DocxConverter


SUPPORTED_FORMATS = {
    'pdf': {'mime': 'application/pdf', 'ext': '.pdf'},
    'png': {'mime': 'image/png', 'ext': '.png'},
    'jpeg': {'mime': 'image/jpeg', 'ext': '.jpg'},
    'webp': {'mime': 'image/webp', 'ext': '.webp'},
    'ico': {'mime': 'image/x-icon', 'ext': '.ico'},
    'tiff': {'mime': 'image/tiff', 'ext': '.tiff'},
}

# Поддерживаемые входные форматы
INPUT_FORMATS = {
    'image': ['.png', '.jpg', '.jpeg'],
    'pdf': ['.pdf'],
    'docx': ['.docx'],
}


def convert_image(image_path: str, output_format: str) -> tuple[bytes, str, str]:
    """
    Конвертирует изображение в указанный формат.
    """
    if output_format not in SUPPORTED_FORMATS:
        raise ValueError(f"Неподдерживаемый формат: {output_format}")

    format_info = SUPPORTED_FORMATS[output_format]
    img = Image.open(image_path)
    buffer = io.BytesIO()

    if output_format == 'pdf':
        img_rgb = img.convert('RGB')
        width, height = img.size
        pdf = canvas.Canvas(buffer, pagesize=(width, height))
        img_buffer = io.BytesIO()
        img_rgb.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        pdf.drawImage(ImageReader(img_buffer), 0, 0, width=width, height=height)
        pdf.save()

    elif output_format == 'jpeg':
        img = img.convert('RGB')
        img.save(buffer, "JPEG", quality=95, optimize=True)

    elif output_format == 'webp':
        img.save(buffer, "WEBP", quality=95, method=6)

    elif output_format == 'ico':
        img = img.convert('RGB')
        img.save(buffer, "ICO")

    elif output_format == 'tiff':
        img.save(buffer, "TIFF", compression='tiff_lzw')

    else:
        img.save(buffer, "PNG", optimize=True)

    buffer.seek(0)
    return buffer.getvalue(), format_info['mime'], format_info['ext']


def convert_pdf_to_image(pdf_path: str, output_format: str, page_number: int = 0) -> tuple[bytes, str, str]:
    """
    Конвертирует PDF в изображение (первую страницу или указанную).
    """
    if output_format not in ['png', 'jpeg']:
        raise ValueError(f"PDF можно конвертировать только в PNG или JPEG")

    doc = fitz.open(pdf_path)
    
    if page_number >= len(doc):
        page_number = 0
    
    page = doc[page_number]
    zoom = 2.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat)
    img_data = pix.tobytes(output_format.upper())
    doc.close()
    
    mime_type = 'image/png' if output_format == 'png' else 'image/jpeg'
    ext = '.png' if output_format == 'png' else '.jpg'
    
    return img_data, mime_type, ext


def convert_pdf_to_docx(pdf_path: str) -> tuple[bytes, str, str]:
    """
    Конвертирует PDF в DOCX (Word документ).
    """
    buffer = io.BytesIO()
    
    # Конвертируем PDF в DOCX
    cv = PDF2DocxConverter(pdf_path)
    cv.convert(buffer)
    cv.close()
    
    buffer.seek(0)
    return buffer.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx'


def convert_docx_to_pdf(docx_path: str) -> tuple[bytes, str, str]:
    """
    Конвертирует DOCX в PDF через reportlab.
    """
    doc = Document(docx_path)
    buffer = io.BytesIO()
    pdf_doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # Добавляем параграфы
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            p = Paragraph(paragraph.text, styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 6))

    # Добавляем таблицы
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                table_data.append(row_data)
        
        if table_data:
            t = Table(table_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), '#EEEEEE'),
                ('GRID', (0, 0), (-1, -1), 1, '#000000'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(t)
            story.append(Spacer(1, 12))

    pdf_doc.build(story)
    buffer.seek(0)
    return buffer.getvalue(), 'application/pdf', '.pdf'


def convert_docx_to_image(docx_path: str, output_format: str) -> tuple[bytes, str, str]:
    """
    Конвертирует DOCX в изображение (через PDF).
    Сначала DOCX → PDF, затем PDF → изображение.
    """
    if output_format not in ['png', 'jpeg']:
        raise ValueError(f"DOCX можно конвертировать только в PNG или JPEG")
    
    # Сначала конвертируем в PDF
    pdf_bytes, _, _ = convert_docx_to_pdf(docx_path)
    
    # Сохраняем PDF во временный файл для PyMuPDF
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
        tmp_pdf.write(pdf_bytes)
        tmp_pdf_path = tmp_pdf.name
    
    try:
        # Конвертируем PDF в изображение
        img_data, mime_type, ext = convert_pdf_to_image(tmp_pdf_path, output_format, page_number=0)
    finally:
        import os
        os.remove(tmp_pdf_path)
    
    return img_data, mime_type, ext


def convert_file(input_path: str, input_type: str, output_format: str) -> tuple[bytes, str, str]:
    """
    Универсальная функция конвертации.
    """
    if input_type == 'image':
        return convert_image(input_path, output_format)
    elif input_type == 'pdf':
        if output_format == 'docx':
            return convert_pdf_to_docx(input_path)
        elif output_format in ['png', 'jpeg']:
            return convert_pdf_to_image(input_path, output_format)
        else:
            raise ValueError(f"PDF можно конвертировать в PNG, JPEG или DOCX")
    elif input_type == 'docx':
        if output_format == 'pdf':
            return convert_docx_to_pdf(input_path)
        elif output_format in ['png', 'jpeg']:
            return convert_docx_to_image(input_path, output_format)
        else:
            raise ValueError(f"DOCX можно конвертировать в PDF, PNG или JPEG")
    else:
        raise ValueError(f"Неподдерживаемый тип файла: {input_type}")


def get_supported_formats() -> list[dict]:
    """Возвращает список поддерживаемых форматов для изображений."""
    return [
        {'id': 'pdf', 'name': 'PDF', 'description': 'Документ PDF'},
        {'id': 'png', 'name': 'PNG', 'description': 'Изображение PNG'},
        {'id': 'jpeg', 'name': 'JPEG', 'description': 'Изображение JPG'},
        {'id': 'webp', 'name': 'WebP', 'description': 'Изображение WebP'},
        {'id': 'ico', 'name': 'ICO', 'description': 'Иконка'},
        {'id': 'tiff', 'name': 'TIFF', 'description': 'Изображение TIFF'},
    ]


def get_input_format_description(filename: str) -> str:
    """Определяет тип файла по расширению."""
    ext = '.' + filename.split('.')[-1].lower()

    if ext in INPUT_FORMATS['image']:
        return 'image'
    elif ext in INPUT_FORMATS['pdf']:
        return 'pdf'
    elif ext in INPUT_FORMATS['docx']:
        return 'docx'
    else:
        return 'unknown'


def get_available_output_formats(input_type: str) -> list[dict]:
    """Возвращает доступные форматы для данного типа входа."""
    if input_type == 'image':
        return get_supported_formats()
    elif input_type == 'pdf':
        return [
            {'id': 'png', 'name': 'PNG', 'description': 'Изображение PNG'},
            {'id': 'jpeg', 'name': 'JPEG', 'description': 'Изображение JPG'},
            {'id': 'docx', 'name': 'DOCX', 'description': 'Документ Word'},
        ]
    elif input_type == 'docx':
        return [
            {'id': 'pdf', 'name': 'PDF', 'description': 'Документ PDF'},
            {'id': 'png', 'name': 'PNG', 'description': 'Изображение PNG'},
            {'id': 'jpeg', 'name': 'JPEG', 'description': 'Изображение JPG'},
        ]
    else:
        return []
