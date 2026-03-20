"""
Конвертация DOCX в PDF и изображения.
"""
import io
import os
import tempfile
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from .config import FONT_ROBOTO, PDF_FONT_SIZE, PDF_LINE_HEIGHT
from .pdf_converter import convert_pdf_to_image


def convert_docx_to_pdf(docx_path: str) -> tuple[bytes, str, str]:
    """
    Конвертирует DOCX в PDF с поддержкой кириллицы.

    Args:
        docx_path: Путь к DOCX файлу

    Returns:
        Кортеж (байты файла, MIME тип, расширение)
    """
    doc = Document(docx_path)
    buffer = io.BytesIO()
    pdf_doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # Регистрируем шрифт Roboto для кириллицы
    cyrillic_style = _get_cyrillic_style(styles)

    # Добавляем параграфы
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            p = Paragraph(paragraph.text, cyrillic_style)
            story.append(p)
            story.append(Spacer(1, 6))

    # Добавляем таблицы
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                table_data.append(row_data)

        if table_data:
            t = Table(table_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), '#EEEEEE'),
                ('GRID', (0, 0), (-1, -1), 1, '#000000'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(t)
            story.append(Spacer(1, 12))

    pdf_doc.build(story)
    buffer.seek(0)
    return buffer.getvalue(), 'application/pdf', '.pdf'


def convert_docx_to_image(docx_path: str, output_format: str) -> tuple[bytes, str, str]:
    """
    Конвертирует DOCX в изображение (через PDF).

    Args:
        docx_path: Путь к DOCX файлу
        output_format: Целевой формат ('png' или 'jpeg')

    Returns:
        Кортеж (байты файла, MIME тип, расширение)

    Raises:
        ValueError: Если формат не поддерживается
    """
    if output_format not in ['png', 'jpeg']:
        raise ValueError(f"DOCX можно конвертировать только в PNG или JPEG")

    # Сначала конвертируем в PDF
    pdf_bytes, _, _ = convert_docx_to_pdf(docx_path)

    # Сохраняем PDF во временный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
        tmp_pdf.write(pdf_bytes)
        tmp_pdf_path = tmp_pdf.name

    try:
        # Конвертируем PDF в изображение
        img_data, mime_type, ext = convert_pdf_to_image(tmp_pdf_path, output_format, page_number=0)
    finally:
        os.remove(tmp_pdf_path)

    return img_data, mime_type, ext


def _get_cyrillic_style(styles) -> ParagraphStyle:
    """
    Получает стиль параграфа с поддержкой кириллицы.

    Args:
        styles: Менеджер стилей reportlab

    Returns:
        Стиль параграфа с кириллическим шрифтом
    """
    try:
        if os.path.exists(FONT_ROBOTO):
            pdfmetrics.registerFont(TTFont('Roboto', FONT_ROBOTO))
            return ParagraphStyle(
                'CyrillicNormal',
                parent=styles['Normal'],
                fontName='Roboto',
                fontSize=PDF_FONT_SIZE,
                leading=PDF_LINE_HEIGHT
            )
    except Exception as e:
        print(f"Warning: Font registration error: {e}")

    # Fallback на стандартный стиль
    return styles['Normal']
